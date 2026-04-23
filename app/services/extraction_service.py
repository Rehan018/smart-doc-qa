from typing import Optional

from pypdf import PdfReader
from docx import Document as DocxDocument


class ExtractedDocument:
    def __init__(self, text: str, metadata: Optional[dict] = None):
        self.text = text
        self.metadata = metadata or {}


class ExtractionService:
    def extract(self, file_path: str, file_type: str) -> ExtractedDocument:
        if file_type == "pdf":
            return self._extract_pdf(file_path)
        elif file_type == "docx":
            return self._extract_docx(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_type}")

    def _extract_pdf(self, file_path: str) -> ExtractedDocument:
        reader = PdfReader(file_path)

        texts = []
        page_map = []

        for i, page in enumerate(reader.pages):
            page_text = page.extract_text() or ""
            page_text = page_text.strip()

            if page_text:
                texts.append(page_text)
                page_map.append({
                    "page": i + 1,
                    "length": len(page_text)
                })

        full_text = "\n\n".join(texts)

        return ExtractedDocument(
            text=self._normalize_text(full_text),
            metadata={"pages": page_map}
        )

    def _extract_docx(self, file_path: str) -> ExtractedDocument:
        doc = DocxDocument(file_path)

        texts = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                texts.append(text)

        full_text = "\n\n".join(texts)

        return ExtractedDocument(
            text=self._normalize_text(full_text),
            metadata={}
        )

    def _normalize_text(self, text: str) -> str:
        # remove excessive whitespace
        lines = [line.strip() for line in text.splitlines()]
        lines = [line for line in lines if line]

        return "\n".join(lines)