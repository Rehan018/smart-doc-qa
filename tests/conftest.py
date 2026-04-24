import os
import tempfile
from pathlib import Path
import sys

import pytest
from fastapi.testclient import TestClient

PROJECT_ROOT = Path(__file__).resolve().parents[1]
if str(PROJECT_ROOT) not in sys.path:
    sys.path.insert(0, str(PROJECT_ROOT))

from app.main import app


@pytest.fixture()
def client():
    return TestClient(app)


@pytest.fixture()
def sample_docx_file():
    from docx import Document

    temp_dir = tempfile.mkdtemp()
    file_path = os.path.join(temp_dir, "sample.docx")

    doc = Document()
    doc.add_heading("Employee Policy", level=1)
    doc.add_paragraph("Employees are entitled to 12 casual leave days per year.")
    doc.add_paragraph("Maternity leave is available as per company policy.")
    doc.save(file_path)

    return file_path
